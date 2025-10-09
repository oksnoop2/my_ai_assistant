# in ./ingestion-service/service.py

import os
import sys
import base64
import logging
import requests
import magic
import io
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from tqdm import tqdm
import trafilatura
from bs4 import BeautifulSoup, XMLParsedAsHTMLWarning
from typing import List, Dict, Any, Optional
from pathlib import Path
from multiprocessing import Pool, cpu_count
import warnings

# --- LlamaIndex for advanced chunking ---
from llama_index.core import Document as LlamaDocument
from llama_index.core.node_parser import SentenceSplitter

# --- Document Parsers ---
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from PIL import Image, UnidentifiedImageError

# --- Suppress specific warnings ---
warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)
warnings.filterwarnings("ignore", category=UserWarning, module='torch')

# --- Logging and Configuration ---
logging.basicConfig(stream=sys.stdout, level=logging.INFO, format="%(asctime)s [%(levelname)s] [%(processName)s] %(message)s")
logger = logging.getLogger(__name__)

try:
    RESOURCE_MANAGER_URL = os.environ["RESOURCE_MANAGER_URL"]
    RAG_SERVICE_URL = os.environ["RAG_SERVICE_URL"]
    LLM_SERVICE_URL = os.environ["LLM_SERVICE_URL"]  # <-- ADDED: This was missing
except KeyError as e:
    logger.error(f"🔥 Critical environment variable missing: {e}")
    sys.exit(1)

# --- Model & Processing Configuration ---
VISION_PROMPT = "Describe this image in detail for a retrieval-augmented generation (RAG) database. Focus on the key subjects, objects, and any text present. Explain charts or diagrams. This will be used as a knowledge base."
BATCH_SIZE = 32  # How many documents to send to the RAG service at once
TEXT_CHUNK_SPLITTER = SentenceSplitter(chunk_size=1024, chunk_overlap=100)

# --- FastAPI App ---
app = FastAPI(title="High-Throughput Data Ingestion Service")

class IngestRequest(BaseModel):
    path: str

# ==============================================================================
# CORE FILE PROCESSING LOGIC
# ==============================================================================

def chunk_text(text: str, metadata: dict) -> List[Dict[str, Any]]:
    """Chunks text using SentenceSplitter and returns a list of document dicts."""
    if not text or not text.strip():
        return []
    
    doc = LlamaDocument(text=text, metadata=metadata)
    nodes = TEXT_CHUNK_SPLITTER.get_nodes_from_documents([doc])
    
    return [
        {"text": node.get_content(), "metadata": {**node.metadata, "chunk_id": i}}
        for i, node in enumerate(nodes)
    ]

def process_file_worker(args) -> List[Dict[str, Any]]:
    """
    A single worker function designed to be called by a multiprocessing Pool.
    It processes one file and returns a list of document chunks.
    """
    file_path_str, root_dir_str = args
    file_path = Path(file_path_str)
    root_dir = Path(root_dir_str)
    
    try:
        relative_path = file_path.relative_to(root_dir).as_posix()
        mime_type = magic.from_file(file_path_str, mime=True)
        metadata = {"source": relative_path, "type": mime_type}
        
        # --- Image Processing ---
        if mime_type.startswith('image/'):
            try:
                with Image.open(file_path) as img:
                    img.verify() # Check for corruption
                
                logger.info(f"🖼️  Processing image via LLM Service: {relative_path}")
            
                with open(file_path, "rb") as image_file:
                    base64_image = base64.b64encode(image_file.read()).decode('utf-8')
                
                # Call the vision model through the LLM service
                payload = {
                    "model_name": "vision", 
                    "prompt": VISION_PROMPT, 
                    "image_base64": base64_image
                }
                response = requests.post(f"{LLM_SERVICE_URL}/completion", json=payload, timeout=300)
                response.raise_for_status()
                description = response.json().get("content", "")

                return [{"text": f"Image Content: {description.strip()}", "metadata": metadata}]
                
            except UnidentifiedImageError:
                logger.warning(f"⚠️ Skipping corrupted or unidentified image: {relative_path}")
                return []
            except Exception as e:
                logger.error(f"❌ Failed to process image {relative_path}: {e}")
                return []

        # --- Text-based File Processing ---
        else:
            raw_text = ""
            logger.info(f"📄 Processing text-based file: {relative_path}")
            with open(file_path, 'rb') as f:
                content = f.read()

            if mime_type == 'application/pdf':
                reader = PdfReader(io.BytesIO(content))
                raw_text = "\n".join(p.extract_text() or "" for p in reader.pages)
            
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                doc = DocxDocument(io.BytesIO(content))
                raw_text = "\n".join(p.text for p in doc.paragraphs)

            elif mime_type.startswith('text/'):
                decoded_content = content.decode('utf-8', errors='ignore')
                if mime_type == 'text/html':
                    extracted_text = trafilatura.extract(decoded_content, include_tables=True)
                    soup = BeautifulSoup(extracted_text or decoded_content, "lxml")
                    raw_text = soup.get_text(separator=" ", strip=True)
                else: # plain text, md, json, etc.
                    raw_text = decoded_content
            
            return chunk_text(raw_text, metadata)

    except Exception as e:
        logger.error(f"❌ Worker failed to process file {file_path_str}: {e}")
        return []
    
    return []


# ==============================================================================
# API ENDPOINTS
# ==============================================================================

@app.on_event("startup")
def startup_event():
    """On startup, verify connectivity to required services."""
    logger.info("🚀 Ingestion service starting up...")
    logger.info("✅ Ingestion service ready.")

@app.get("/health")
def health():
    return {"status": "ok"}

@app.post("/ingest")
def ingest_data(payload: IngestRequest):
    """The main endpoint to start the parallelized ingestion process."""
    ingest_path = Path(payload.path)
    if not ingest_path.is_dir():
        raise HTTPException(status_code=404, detail=f"Directory not found: {ingest_path}")

    logger.info(f"🚀 Starting ingestion for directory: {ingest_path}")

    files_to_process = [str(p) for p in ingest_path.rglob("*") if p.is_file()]
    if not files_to_process:
        return {"status": "complete", "message": "No files found to process."}

    all_docs_to_load = []
    num_workers = min(cpu_count(), len(files_to_process))
    logger.info(f"Found {len(files_to_process)} files. Starting parallel processing with {num_workers} workers.")
    
    # Create arguments for the worker pool
    worker_args = [(file_path, str(ingest_path)) for file_path in files_to_process]

    with Pool(processes=num_workers) as pool:
        # Use tqdm to show progress for the file processing
        results = list(tqdm(pool.imap(process_file_worker, worker_args), total=len(files_to_process), desc="Processing files"))

    # Flatten the list of lists into a single list of document chunks
    all_docs_to_load = [doc for sublist in results if sublist for doc in sublist]

    logger.info(f"✅ File processing complete. Generated {len(all_docs_to_load)} document chunks.")
    if not all_docs_to_load:
        return {"status": "complete", "message": "Processing finished, but no usable document chunks were generated."}

    logger.info("📦 Now loading documents into RAG service in batches...")
    total_loaded = 0
    for i in tqdm(range(0, len(all_docs_to_load), BATCH_SIZE), desc="Loading to RAG"):
        batch = all_docs_to_load[i:i + BATCH_SIZE]
        try:
            response = requests.post(f"{RAG_SERVICE_URL}/add_document_batch", json={"documents": batch}, timeout=180)
            response.raise_for_status()
            total_loaded += len(batch)
        except requests.RequestException as e:
            logger.error(f"❌ Failed to send batch to RAG service: {e}")
            logger.error(f"Failed batch metadata: {[doc.get('metadata') for doc in batch]}")
    
    final_message = f"Ingestion complete. Successfully loaded {total_loaded}/{len(all_docs_to_load)} documents into the knowledge graph."
    logger.info(f"🎉 {final_message}")
    return {"status": "complete", "message": final_message}
